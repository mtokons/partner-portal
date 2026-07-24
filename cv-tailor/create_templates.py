"""
One-time script: generate the three starter DOCX template files.
Run inside the container or locally: python create_templates.py
The templates are simple placeholder documents; replace them with your
branded DOCX files and the /generate endpoint will use them automatically.
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

TEMPLATES_DIR = Path(__file__).parent / "templates"
TEMPLATES_DIR.mkdir(exist_ok=True)

COLOR_GIZ  = RGBColor(0x00, 0x5F, 0xA0)   # GIZ blue
COLOR_EU   = RGBColor(0x00, 0x3D, 0x99)   # EU blue
COLOR_UCEP = RGBColor(0x00, 0x91, 0x5E)   # UN green

def make_template(path: Path, title: str, subtitle: str, accent: RGBColor):
    doc = Document()
    # Header branding
    for sec in doc.sections:
        sec.top_margin    = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin   = Inches(1.25)
        sec.right_margin  = Inches(1.25)

    h = doc.add_heading(f"[EXPERT_NAME] – {title}", 0)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in h.runs:
        run.font.color.rgb = accent
        run.font.size = Pt(16)

    sub = doc.add_paragraph(subtitle)
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = accent
    sub.runs[0].font.size = Pt(10)

    doc.add_paragraph("")
    doc.add_paragraph("[TOR_MATCH_PCT]% TOR Match Rating").runs[0].bold = True
    doc.add_paragraph("")

    doc.save(str(path))
    print(f"Created {path.name}")

make_template(TEMPLATES_DIR / "giz.docx",  "Curriculum Vitae",     "GIZ Corporate Format",        COLOR_GIZ)
make_template(TEMPLATES_DIR / "eu.docx",   "Curriculum Vitae",     "EU Standard Europass Format", COLOR_EU)
make_template(TEMPLATES_DIR / "ucep.docx", "Competency Profile",   "UCEP / UN Agency Format",     COLOR_UCEP)
print("Done.")

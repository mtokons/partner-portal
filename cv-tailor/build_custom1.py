"""
Custom CV Format 1 — replicates the exact layout of the GIZ-style CV used in
the SCCG Project Partner pipeline (based on CV_Arifin.docx analysis):
  • A4 portrait  (most sections)
  • Arial 7.5 pt throughout
  • Blue header rows (#C6D9F1) in tables
  • Personal-info block with tab-aligned labels
  • Proposed role row  (1×2 table)
  • Education table    (N×4)
  • Training table     (N×6)
  • Language table     (N×4)
  • Key qualifications (bullet list)
  • Regional experience table (N×4)
  • Professional experience table (N×7)  — landscape page
  • Publications / other information (bullet list)

All person data is passed in via the `person` dict; AI content comes from the
standard TailorResult JSON.
"""
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
import io

FONT_NAME    = "Arial"
FONT_SIZE    = Pt(7.5)
HDR_BG       = "C6D9F1"   # light blue header fill


def _set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _cell_run(cell, text: str, bold=False, size: Pt = FONT_SIZE, color: str | None = None, clear=True):
    """Set text in a table cell with exact formatting."""
    if clear:
        for para in cell.paragraphs:
            for run in para.runs:
                run.text = ""
        # Use first paragraph
        para = cell.paragraphs[0]
    else:
        para = cell.paragraphs[-1]

    para.clear()
    run = para.add_run(text)
    run.font.name  = FONT_NAME
    run.font.size  = size
    run.font.bold  = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def _add_run(para, text: str, bold=False, size: Pt = FONT_SIZE):
    run = para.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    return run


def _header_row(table, headers: list[str]):
    """Set up the first row of a table as a blue header."""
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        _set_cell_bg(cell, HDR_BG)
        _cell_run(cell, h, bold=True, color="000000")


def _add_label_value_para(doc, label: str, value: str):
    """Add a tab-separated label: value paragraph."""
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    r = p.add_run(label + ":")
    r.font.name, r.font.size = FONT_NAME, FONT_SIZE
    r2 = p.add_run("\t\t" + value)
    r2.font.name, r2.font.size = FONT_NAME, FONT_SIZE
    return p


def _section_heading(doc, text: str):
    """Add a plain section heading paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    r = p.add_run(text)
    r.font.name, r.font.size, r.font.bold = FONT_NAME, FONT_SIZE, False
    return p


def _bullet(doc, text: str, bold_prefix: str = ""):
    """Add a bullet point paragraph."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        rb = p.add_run(bold_prefix)
        rb.font.name, rb.font.size, rb.font.bold = FONT_NAME, FONT_SIZE, True
        rv = p.add_run(text)
        rv.font.name, rv.font.size = FONT_NAME, FONT_SIZE
    else:
        r = p.add_run(text)
        r.font.name, r.font.size = FONT_NAME, FONT_SIZE
    return p


def _ensure_list_bullet_style(doc):
    """Add a simple List Bullet style if it doesn't exist."""
    if "List Bullet" not in [s.name for s in doc.styles]:
        from docx.oxml.ns import nsmap
        style = doc.styles.add_style("List Bullet", 1)  # PARAGRAPH
        style.font.name = FONT_NAME
        style.font.size = FONT_SIZE


def build_custom1_docx(result: dict, person: dict) -> bytes:
    """
    Build a Custom CV Format 1 DOCX from:
      result   — standard TailorResult JSON (expert_name, sections, matrix_matches)
      person   — personal data dict with keys:
                 proposedRole, familyName, firstName, dateOfBirth,
                 nationality, placeOfResidence, email, tel,
                 membership, otherSkills, presentPosition,
                 education  = [{institution, dateFrom, dateTo, degree}]
                 training   = [{course, provider, location, year, competency, certificate}]
                 languages  = [{language, reading, speaking, writing}]
                 regions    = [{nr, region, country, dates}]
                 experience = [{dateFrom, dateTo, wd, location, company, position, description}]
                 publications = [str]
    """
    doc = Document()
    _ensure_list_bullet_style(doc)

    # ── Page setup: A4 portrait ───────────────────────────────────────────────
    for sec in doc.sections:
        sec.page_width    = Inches(8.27)
        sec.page_height   = Inches(11.69)
        sec.top_margin    = Inches(0.709)
        sec.bottom_margin = Inches(0.709)
        sec.left_margin   = Inches(0.394)
        sec.right_margin  = Inches(0.492)

    # ── CURRICULUM VITAE heading ──────────────────────────────────────────────
    h = doc.add_heading("CURRICULUM VITAE", level=0)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.name = FONT_NAME
        run.font.size = FONT_SIZE

    doc.add_paragraph()

    # ── Personal info ─────────────────────────────────────────────────────────
    family_name   = person.get("familyName") or result.get("expert_name", "").split()[-1] if result.get("expert_name") else ""
    first_names   = person.get("firstName") or (result.get("expert_name", "").split()[0] if result.get("expert_name") else "")
    _add_label_value_para(doc, "Family name",        family_name)
    _add_label_value_para(doc, "First names",        first_names)
    _add_label_value_para(doc, "Date of Birth",      person.get("dateOfBirth", ""))
    _add_label_value_para(doc, "Nationality",        person.get("nationality", ""))
    _add_label_value_para(doc, "Place of residence", person.get("placeOfResidence", ""))
    _add_label_value_para(doc, "Email",              person.get("email", ""))
    _add_label_value_para(doc, "Tel.",               person.get("tel", ""))

    doc.add_paragraph()

    # ── Proposed role ─────────────────────────────────────────────────────────
    _section_heading(doc, "")
    role_tbl = doc.add_table(rows=1, cols=2)
    role_tbl.style = "Table Grid"
    _cell_run(role_tbl.rows[0].cells[0], " Proposed role in the project:")
    _cell_run(role_tbl.rows[0].cells[1], " " + person.get("proposedRole", result.get("expert_name", "")))
    doc.add_paragraph()

    # ── Education ─────────────────────────────────────────────────────────────
    _section_heading(doc, "Education:")
    education = person.get("education") or []
    edu_rows  = max(2, len(education) + 1)
    edu_tbl   = doc.add_table(rows=edu_rows, cols=4)
    edu_tbl.style = "Table Grid"
    _header_row(edu_tbl, ["Institution", "Date from", "Date to", "Degree(s) or Diploma(s) obtained:"])
    for i, entry in enumerate(education):
        row = edu_tbl.rows[i + 1]
        _cell_run(row.cells[0], entry.get("institution", ""))
        _cell_run(row.cells[1], entry.get("dateFrom", ""))
        _cell_run(row.cells[2], entry.get("dateTo", ""))
        _cell_run(row.cells[3], entry.get("degree", ""))
    doc.add_paragraph()

    # ── Training ──────────────────────────────────────────────────────────────
    _section_heading(doc, "Training:")
    training  = person.get("training") or []
    train_rows = max(2, len(training) + 1)
    train_tbl  = doc.add_table(rows=train_rows, cols=6)
    train_tbl.style = "Table Grid"
    _header_row(train_tbl, [
        "Training / Professional Development",
        "Provider", "Location", "Year",
        "Key Competency Gained",
        "Certificate / Qualification(s) obtained:",
    ])
    for i, entry in enumerate(training):
        row = train_tbl.rows[i + 1]
        _cell_run(row.cells[0], entry.get("course", ""), bold=True)
        _cell_run(row.cells[1], entry.get("provider", ""))
        _cell_run(row.cells[2], entry.get("location", ""))
        _cell_run(row.cells[3], entry.get("year", ""))
        _cell_run(row.cells[4], entry.get("competency", ""))
        _cell_run(row.cells[5], entry.get("certificate", ""))
    doc.add_paragraph()

    # ── Language skills ───────────────────────────────────────────────────────
    _section_heading(doc, "Language skills: Indicate competence from A1 (beginner) to C2 (proficient).")
    languages  = person.get("languages") or []
    lang_rows  = max(2, len(languages) + 1)
    lang_tbl   = doc.add_table(rows=lang_rows, cols=4)
    lang_tbl.style = "Table Grid"
    _header_row(lang_tbl, ["Language", "Reading", "Speaking", "Writing"])
    for i, entry in enumerate(languages):
        row = lang_tbl.rows[i + 1]
        _cell_run(row.cells[0], entry.get("language", ""))
        lvl = entry.get("level") or entry.get("reading") or ""
        _cell_run(row.cells[1], entry.get("reading", lvl))
        _cell_run(row.cells[2], entry.get("speaking", lvl))
        _cell_run(row.cells[3], entry.get("writing", lvl))
    doc.add_paragraph()

    # ── Membership ────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    _add_run(p, "Membership of professional bodies: ", bold=True)
    _add_run(p, person.get("membership", ""))

    # ── Other skills ──────────────────────────────────────────────────────────
    p2 = doc.add_paragraph()
    _add_run(p2, "Other skills (e.g. computer literacy, etc.): ")
    _add_run(p2, person.get("otherSkills", ""))

    # ── Present position ──────────────────────────────────────────────────────
    p3 = doc.add_paragraph()
    _add_run(p3, "Present position: ")
    _add_run(p3, person.get("presentPosition", ""))

    doc.add_paragraph()

    # ── Key qualifications ────────────────────────────────────────────────────
    _section_heading(doc, "Key qualifications (relevant to the project):")
    doc.add_paragraph()
    # Pull from AI tailored sections
    for sec in result.get("sections", []):
        if sec.get("tailored"):
            _bullet(doc, sec["tailored"])

    doc.add_paragraph()

    # ── Specific regional experience ──────────────────────────────────────────
    _section_heading(doc, "Specific experience in the region:")
    regions   = person.get("regions") or []
    reg_rows  = max(2, len(regions) + 1)
    reg_tbl   = doc.add_table(rows=reg_rows, cols=4)
    reg_tbl.style = "Table Grid"
    _header_row(reg_tbl, ["Nr.", "Region", "Country", "Date from - Date to"])
    for i, entry in enumerate(regions):
        row = reg_tbl.rows[i + 1]
        _cell_run(row.cells[0], entry.get("nr", str(i + 1) + "."), bold=True)
        _cell_run(row.cells[1], entry.get("region", ""), bold=True)
        _cell_run(row.cells[2], entry.get("country", ""), bold=True)
        _cell_run(row.cells[3], entry.get("dates", ""))
    doc.add_paragraph()

    # ── Professional experience — landscape page ──────────────────────────────
    # Add a section break + change orientation to landscape
    new_sec_para = doc.add_paragraph()
    new_sec_para.add_run()
    from docx.oxml import OxmlElement as _OE
    pPr = new_sec_para._p.get_or_add_pPr()
    sectPr = _OE("w:sectPr")
    pgSz = _OE("w:pgSz")
    pgSz.set(qn("w:w"), str(int(11.69 * 1440)))   # landscape: width = A4 height
    pgSz.set(qn("w:h"), str(int(8.27  * 1440)))
    pgSz.set(qn("w:orient"), "landscape")
    pgMar = _OE("w:pgMar")
    pgMar.set(qn("w:top"),    str(int(0.394 * 1440)))
    pgMar.set(qn("w:bottom"), str(int(0.492 * 1440)))
    pgMar.set(qn("w:left"),   str(int(0.709 * 1440)))
    pgMar.set(qn("w:right"),  str(int(0.709 * 1440)))
    sectPr.append(pgSz)
    sectPr.append(pgMar)
    pPr.append(sectPr)

    _section_heading(doc, "Professional experience:")
    experience  = person.get("experience") or []
    exp_rows    = max(2, len(experience) + 1)
    exp_tbl     = doc.add_table(rows=exp_rows, cols=7)
    exp_tbl.style = "Table Grid"
    _header_row(exp_tbl, ["Date from", "Date to", "WD", "Location", "Company", "Position", "Description"])
    for i, entry in enumerate(experience):
        row = exp_tbl.rows[i + 1]
        _cell_run(row.cells[0], entry.get("dateFrom", ""))
        _cell_run(row.cells[1], entry.get("dateTo", ""))
        _cell_run(row.cells[2], entry.get("wd", ""))
        _cell_run(row.cells[3], entry.get("location", ""))
        _cell_run(row.cells[4], entry.get("company", ""), bold=True)
        _cell_run(row.cells[5], entry.get("position", ""), bold=True)
        # Description cell: Project Title + Donor + main description
        desc_cell = row.cells[6]
        desc_cell.paragraphs[0].clear()
        desc_para = desc_cell.paragraphs[0]
        if entry.get("projectTitle"):
            rb = desc_para.add_run("Project Title: ")
            rb.font.name, rb.font.size, rb.font.bold = FONT_NAME, FONT_SIZE, True
            rv = desc_para.add_run(entry["projectTitle"])
            rv.font.name, rv.font.size = FONT_NAME, FONT_SIZE
        if entry.get("donor"):
            desc_para.add_run("\n")
            rd = desc_para.add_run("Donor: ")
            rd.font.name, rd.font.size, rd.font.bold = FONT_NAME, FONT_SIZE, True
            rdv = desc_para.add_run(entry["donor"])
            rdv.font.name, rdv.font.size = FONT_NAME, FONT_SIZE
        if entry.get("description"):
            desc_para.add_run("\n")
            rdesc = desc_para.add_run(entry.get("description", ""))
            rdesc.font.name, rdesc.font.size = FONT_NAME, FONT_SIZE

    doc.add_paragraph()

    # ── Publications / Other information ─────────────────────────────────────
    _section_heading(doc, "Other relevant information (e.g. publications):")
    for pub in (person.get("publications") or []):
        _bullet(doc, pub)

    # Also append matrix evidence as reference
    if result.get("matrix_matches"):
        doc.add_paragraph()
        _section_heading(doc, "Evaluation Matrix Alignment:")
        mat_tbl = doc.add_table(rows=1, cols=4)
        mat_tbl.style = "Table Grid"
        _header_row(mat_tbl, ["Requirement", "Evidence", "Score", "Max"])
        for m in result.get("matrix_matches", []):
            row = mat_tbl.add_row()
            _cell_run(row.cells[0], str(m.get("requirement", ""))[:150])
            _cell_run(row.cells[1], str(m.get("evidence", ""))[:200])
            _cell_run(row.cells[2], str(m.get("score", 0)))
            _cell_run(row.cells[3], str(m.get("max_score", 0)))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Standalone: generate a blank placeholder template ────────────────────────
if __name__ == "__main__":
    TEMPLATES_DIR = Path(__file__).parent / "templates"
    TEMPLATES_DIR.mkdir(exist_ok=True)

    placeholder_result = {
        "expert_name": "EXPERT NAME",
        "tor_match_pct": 0,
        "sections": [{"section": "Key Qualifications", "tailored": "[KEY QUALIFICATIONS WILL BE FILLED BY AI]", "keywords": []}],
        "matrix_matches": [],
    }
    placeholder_person = {
        "proposedRole": "[PROPOSED ROLE]",
        "familyName": "[FAMILY NAME]",
        "firstName": "[FIRST NAMES]",
        "dateOfBirth": "[DATE OF BIRTH]",
        "nationality": "[NATIONALITY]",
        "placeOfResidence": "[PLACE OF RESIDENCE]",
        "email": "[EMAIL]",
        "tel": "[TEL]",
        "membership": "[PROFESSIONAL MEMBERSHIPS]",
        "otherSkills": "[OTHER SKILLS]",
        "presentPosition": "[PRESENT POSITION]",
        "education": [{"institution": "[INSTITUTION]", "dateFrom": "[MM/YYYY]", "dateTo": "[MM/YYYY]", "degree": "[DEGREE / DIPLOMA]"}],
        "training": [{"course": "[TRAINING COURSE]", "provider": "[PROVIDER]", "location": "[LOCATION]", "year": "[YEAR]", "competency": "[COMPETENCY GAINED]", "certificate": "[CERTIFICATE]"}],
        "languages": [{"language": "[LANGUAGE]", "reading": "[LEVEL]", "speaking": "[LEVEL]", "writing": "[LEVEL]"}],
        "regions": [{"nr": "1.", "region": "[REGION]", "country": "[COUNTRY]", "dates": "[MM/YYYY – MM/YYYY]"}],
        "experience": [{"dateFrom": "[MM/YYYY]", "dateTo": "[MM/YYYY]", "wd": "[MONTHS/DAYS]", "location": "[COUNTRY]", "company": "[COMPANY / ORGANISATION]", "position": "[POSITION TITLE]", "projectTitle": "[PROJECT TITLE]", "donor": "[DONOR]", "description": "[MAIN ACTIVITIES AND RESPONSIBILITIES]"}],
        "publications": ["[PUBLICATION OR OTHER RELEVANT INFORMATION]"],
    }
    docx_bytes = build_custom1_docx(placeholder_result, placeholder_person)
    out_path = TEMPLATES_DIR / "custom1.docx"
    out_path.write_bytes(docx_bytes)
    print(f"Created {out_path}")

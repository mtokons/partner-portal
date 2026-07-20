"""
CV Tailoring microservice — FastAPI
Architecture:
  1. Parse CV (PDF / DOCX) → extract text using pypdf / python-docx
  2. Call Gemini Flash with TOR + Evaluation Matrix context → structured JSON
  3. Inject JSON into a target DOCX template (python-docx) → download
"""
import os
import io
import json
import re
import hashlib
import logging
from typing import Optional, Any
from pathlib import Path
import jinja2
import subprocess
import tempfile
import shutil

# pyrefly: ignore [missing-import]
from fastapi import FastAPI, File, Form, UploadFile, HTTPException, Request
# pyrefly: ignore [missing-import]
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
import httpx

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

app = FastAPI(title="CV Tailor API", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")
GEMINI_MODEL   = os.environ.get("GEMINI_MODEL", "gemini-flash-latest")
TEMPLATES_DIR  = Path(__file__).parent / "templates"
CACHE_DIR      = Path(__file__).parent / "cache"
CACHE_DIR.mkdir(exist_ok=True)

# Delimiters are changed so they don't clash with LaTeX brackets
latex_jinja_env = jinja2.Environment(
    block_start_string='((%',
    block_end_string='%))',
    variable_start_string='(((',
    variable_end_string=')))',
    comment_start_string='((#',
    comment_end_string='#))',
    loader=jinja2.FileSystemLoader(str(TEMPLATES_DIR))
)

def escape_latex(text: str) -> str:
    """Escapes special LaTeX characters in plain text to prevent compile errors."""
    if not text:
        return ""
    latex_chars = {
        '&': r'\&',
        '%': r'\%',
        '$': r'\$',
        '#': r'\#',
        '_': r'\_',
        '{': r'\{',
        '}': r'\}',
        '~': r'\textasciitilde{}',
        '^': r'\textasciicircum{}',
        '\\': r'\textbackslash{}',
    }
    regex = re.compile('|'.join(re.escape(key) for key in latex_chars.keys()))
    return regex.sub(lambda match: latex_chars[match.group()], text)

def _escape_latex_recursive(data: Any) -> Any:
    if isinstance(data, str):
        return escape_latex(data)
    elif isinstance(data, list):
        return [_escape_latex_recursive(x) for x in data]
    elif isinstance(data, dict):
        return {k: _escape_latex_recursive(v) for k, v in data.items()}
    return data

def create_placeholder_pdf(message: str) -> bytes:
    # A lightweight, valid minimal PDF file
    pdf = (
        "%PDF-1.4\n"
        "1 0 obj\n<< /Type /Catalog /Pages 2 0 R >>\nendobj\n"
        "2 0 obj\n<< /Type /Pages /Kids [3 0 R] /Count 1 >>\nendobj\n"
        "3 0 obj\n<< /Type /Page /Parent 2 0 R /MediaBox [0 0 595 842] /Resources << /Font << /F1 4 0 R >> >> /Contents 5 0 R >>\nendobj\n"
        "4 0 obj\n<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>\nendobj\n"
        "5 0 obj\n<< /Length 250 >>\nstream\n"
        "BT\n/F1 16 Tf\n50 750 Td\n(LaTeX Compilation Warning) Tj\n"
        "0 -40 Td\n/F1 11 Tf\n"
        f"({message[:80]}) Tj\n"
        "0 -25 Td\n"
        "(To generate the actual compiled PDF, make sure pdflatex is installed on the host) Tj\n"
        "0 -15 Td\n"
        "(or run the service inside the Docker container, which installs texlive.) Tj\n"
        "ET\nendstream\nendobj\n"
        "xref\n0 6\n"
        "0000000000 65535 f \n"
        "0000000009 00000 n \n"
        "0000000056 00000 n \n"
        "0000000111 00000 n \n"
        "0000000212 00000 n \n"
        "0000000282 00000 n \n"
        "trailer\n<< /Size 6 /Root 1 0 R >>\n"
        "startxref\n550\n%%EOF\n"
    )
    return pdf.encode("latin-1")

def _build_pdf(result: dict, person_data: dict | None = None) -> bytes:
    logger.info("Starting LaTeX PDF compilation...")
    # 1. Escape all dynamic strings
    escaped_result = _escape_latex_recursive(result)
    escaped_person_data = _escape_latex_recursive(person_data) if person_data else None
    
    # 2. Render LaTeX template using Jinja2
    try:
        template = latex_jinja_env.get_template("latex_modern.tex")
        rendered_latex = template.render(result=escaped_result, person_data=escaped_person_data)
    except Exception as e:
        logger.error("LaTeX template rendering failed: %s", e)
        return create_placeholder_pdf(f"Template render error: {e}")
        
    # 3. Check for pdflatex availability
    if not shutil.which("pdflatex"):
        logger.warning("pdflatex compiler not found on system path. Using placeholder fallback.")
        return create_placeholder_pdf("pdflatex compiler not found on this machine.")
        
    # 4. Compile to PDF in a temporary directory
    with tempfile.TemporaryDirectory() as tempdir:
        tex_path = os.path.join(tempdir, "cv.tex")
        with open(tex_path, "w", encoding="utf-8") as f:
            f.write(rendered_latex)
            
        try:
            # Run pdflatex twice for table positioning & page numbers
            for run_idx in range(2):
                process = subprocess.run(
                    ["pdflatex", "-interaction=nonstopmode", "cv.tex"],
                    cwd=tempdir,
                    stdout=subprocess.PIPE,
                    stderr=subprocess.PIPE,
                    timeout=20
                )
                if process.returncode != 0 and run_idx == 1:
                    logger.warning("pdflatex returned non-zero code: %d", process.returncode)
                    
            pdf_path = os.path.join(tempdir, "cv.pdf")
            if os.path.exists(pdf_path):
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                logger.info("LaTeX CV compiled successfully (%d bytes).", len(pdf_bytes))
                return pdf_bytes
            else:
                log_path = os.path.join(tempdir, "cv.log")
                log_snippet = ""
                if os.path.exists(log_path):
                    with open(log_path, "r", encoding="utf-8", errors="ignore") as lf:
                        log_snippet = lf.read()[-500:]
                logger.error("pdflatex output PDF not found. Log snippet:\n%s", log_snippet)
                return create_placeholder_pdf(f"pdflatex output PDF not found. compilation error.")
        except subprocess.TimeoutExpired:
            logger.error("pdflatex compilation timed out (20s).")
            return create_placeholder_pdf("pdflatex compilation timed out.")
        except Exception as e:
            logger.exception("pdflatex execution failed")
            return create_placeholder_pdf(f"pdflatex execution failed: {e}")


# ── Output templates available ───────────────────────────────────────────────
TEMPLATE_META = {
    "giz":          {"label": "GIZ Corporate Format",   "description": "Strict multi-page tables, GIZ standard"},
    "eu":           {"label": "Standard EU CV",          "description": "Chronological grid layout"},
    "ucep":         {"label": "UCEP / UN Format",         "description": "UN-style competency profile"},
    "custom1":      {"label": "Custom CV Format 1",       "description": "Exact GIZ-style layout: personal info block, education/training/language/region/experience tables, Arial 7.5pt"},
    "latex_modern": {"label": "LaTeX Modern CV",        "description": "Professional LaTeX compiled PDF"},
}


# ── Pydantic models ───────────────────────────────────────────────────────────
class TailoredSection(BaseModel):
    section:   str
    original:  Optional[str] = None
    tailored:  str
    keywords:  list[str] = []

class MatrixMatch(BaseModel):
    requirement: str
    evidence:    str
    score:       float
    max_score:   float

class TailorResult(BaseModel):
    expert_name:    str
    tor_match_pct:  float
    sections:       list[TailoredSection]
    matrix_matches: list[MatrixMatch]
    provider:       str


# ── Document parsing ──────────────────────────────────────────────────────────
def _extract_text(content: bytes, filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".pdf"):
        try:
            from pypdf import PdfReader
            reader = PdfReader(io.BytesIO(content))
            text = "\n".join(p.extract_text() or "" for p in reader.pages).strip()
            return text
        except Exception as e:
            raise HTTPException(400, f"PDF parse error: {e}")
    if lower.endswith(".docx"):
        try:
            from docx import Document
            doc = Document(io.BytesIO(content))
            parts: list[str] = []
            # 1. Paragraphs
            for p in doc.paragraphs:
                if p.text.strip():
                    parts.append(p.text)
            # 2. Tables — CVs store most data in tables, so this is essential
            for table in doc.tables:
                for row in table.rows:
                    cells = [c.text.strip() for c in row.cells if c.text.strip()]
                    if cells:
                        parts.append(" | ".join(cells))
            return "\n".join(parts).strip()
        except Exception as e:
            raise HTTPException(400, f"DOCX parse error: {e}")
    if lower.endswith((".doc", ".rtf")):
        # legacy formats — best-effort decode
        return content.decode("utf-8", errors="ignore").strip()
    # Plain text fallback
    return content.decode("utf-8", errors="ignore").strip()


IMAGE_EXT = (".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif")

def _is_image(filename: str) -> bool:
    return filename.lower().endswith(IMAGE_EXT)

def _image_mime(filename: str) -> str:
    lower = filename.lower()
    if lower.endswith(".png"): return "image/png"
    if lower.endswith((".jpg", ".jpeg")): return "image/jpeg"
    if lower.endswith(".webp"): return "image/webp"
    if lower.endswith(".gif"): return "image/gif"
    if lower.endswith((".tiff", ".tif")): return "image/tiff"
    if lower.endswith(".bmp"): return "image/bmp"
    return "image/png"


# ── Gemini call (plain REST — same pattern as the Node runtime) ───────────────
async def _call_gemini(prompt: str) -> str:
    if not GEMINI_API_KEY:
        return ""
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [{"text": prompt}]}],
        "generationConfig": {"temperature": 0.2, "responseMimeType": "application/json"},
    }
    async with httpx.AsyncClient(timeout=120) as client:
        r = await client.post(url, json=payload)
        if not r.is_success:
            raise HTTPException(502, f"Gemini error {r.status_code}: {r.text[:200]}")
        data = r.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]


async def _call_gemini_vision(prompt: str, image_bytes: bytes, mime_type: str) -> str:
    """Send an image + prompt to Gemini (vision) — used for scanned docs / photos."""
    if not GEMINI_API_KEY:
        return ""
    import base64
    b64 = base64.b64encode(image_bytes).decode("ascii")
    url = f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}"
    payload = {
        "contents": [{"parts": [
            {"inline_data": {"mime_type": mime_type, "data": b64}},
            {"text": prompt},
        ]}],
        "generationConfig": {"temperature": 0.2, "responseMimeType": "application/json"},
    }
    async with httpx.AsyncClient(timeout=180) as client:
        r = await client.post(url, json=payload)
        if not r.is_success:
            raise HTTPException(502, f"Gemini vision error {r.status_code}: {r.text[:200]}")
        data = r.json()
        return data["candidates"][0]["content"]["parts"][0]["text"]


async def _ocr_image_to_text(image_bytes: bytes, mime_type: str) -> str:
    """Use Gemini vision to transcribe ALL text from an image/scanned document."""
    prompt = (
        "Transcribe ALL text visible in this document image exactly as written, "
        "preserving structure (headings, tables as 'A | B | C' rows, bullet points). "
        "Do not summarise or omit anything. Return JSON: {\"text\": \"<full transcription>\"}"
    )
    raw = await _call_gemini_vision(prompt, image_bytes, mime_type)
    try:
        return _parse_json_loose(raw).get("text", "").strip()
    except Exception:
        return raw.strip()



def _parse_json_loose(text: str) -> dict:
    text = text.strip()
    m = re.search(r"```(?:json)?\s*([\s\S]*?)```", text, re.I)
    if m:
        text = m.group(1).strip()
    start = next((i for i, c in enumerate(text) if c in "{["), 0)
    text = text[start:]
    last = max(text.rfind("}"), text.rfind("]"))
    if last >= 0:
        text = text[: last + 1]
    return json.loads(text)


# ── Mock fallback for keyless environments ────────────────────────────────────
def _mock_tailor(cv_text: str, tor_text: str, criteria: list[dict]) -> dict:
    lines = [l for l in cv_text.split("\n") if len(l) > 30]
    name_m = re.search(r"^([A-Z][a-z]+ [A-Z][a-z]+)", cv_text, re.M)
    return {
        "expert_name": name_m.group(1) if name_m else "Expert",
        "tor_match_pct": 78.0,
        "sections": [
            {
                "section": "Professional Summary",
                "original": lines[0] if lines else cv_text[:200],
                "tailored": f"A highly experienced professional whose background strongly aligns with the TOR requirements. {lines[0] if lines else ''}",
                "keywords": ["TVET", "development", "curriculum"],
            },
            {
                "section": "Key Qualifications",
                "original": None,
                "tailored": "Extensive experience relevant to the stated evaluation matrix criteria, with a track record of measurable impact.",
                "keywords": [],
            },
        ],
        "matrix_matches": [
            {
                "requirement": c.get("label", c.get("text", "Criterion"))[:80],
                "evidence":    f"Evidence found in CV: {lines[min(i, len(lines)-1)] if lines else 'see CV text'}",
                "score":       round(float(c.get("maxPoints", c.get("max_score", 2))) * 0.75, 2),
                "max_score":   float(c.get("maxPoints", c.get("max_score", 2))),
            }
            for i, c in enumerate(criteria[:6])
        ],
        "provider": "mock",
    }


# ── DOCX generator ────────────────────────────────────────────────────────────
def _build_docx(result: dict, template_id: str, person_data: dict | None = None) -> bytes:
    # Custom CV Format 1 — full structured layout
    if template_id == "custom1":
        try:
            from build_custom1 import build_custom1_docx
            return build_custom1_docx(result, person_data or {})
        except Exception as e:
            logger.warning("custom1 build failed (%s), falling back to default", e)
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    tpl_path = TEMPLATES_DIR / f"{template_id}.docx"
    doc = Document(str(tpl_path)) if tpl_path.exists() else Document()

    def heading(text: str, level: int = 1):
        p = doc.add_heading(text, level=level)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        return p

    def body(text: str):
        p = doc.add_paragraph(text)
        p.style.font.size = Pt(10)
        return p

    if not tpl_path.exists():
        # Build a clean default document
        h = doc.add_heading(result.get("expert_name", "Expert"), 0)
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
        match_pct = result.get("tor_match_pct", 0)
        doc.add_paragraph(f"TOR Match Rating: {match_pct:.0f}%")
        doc.add_paragraph("")

    # Sections
    for sec in result.get("sections", []):
        heading(sec["section"], level=1)
        if sec.get("original"):
            orig_para = doc.add_paragraph()
            orig_para.add_run("Original: ").bold = True
            orig_para.add_run(sec["original"][:600])
        body(sec.get("tailored", ""))
        kws = sec.get("keywords", [])
        if kws:
            doc.add_paragraph("Keywords: " + ", ".join(kws))

    # Matrix match table
    if result.get("matrix_matches"):
        doc.add_page_break()
        heading("Evaluation Matrix Alignment", level=1)
        tbl = doc.add_table(rows=1, cols=4)
        tbl.style = "Table Grid"
        hdr = tbl.rows[0].cells
        hdr[0].text = "Requirement"
        hdr[1].text = "Evidence"
        hdr[2].text = "Score"
        hdr[3].text = "Max"
        for m in result.get("matrix_matches", []):
            row = tbl.add_row().cells
            row[0].text = str(m.get("requirement", ""))[:120]
            row[1].text = str(m.get("evidence", ""))[:200]
            row[2].text = str(m.get("score", 0))
            row[3].text = str(m.get("max_score", 0))

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── API routes ────────────────────────────────────────────────────────────────

@app.get("/health")
def health():
    return {"status": "ok", "model": GEMINI_MODEL, "ai_live": bool(GEMINI_API_KEY)}


@app.get("/templates")
def list_templates():
    return TEMPLATE_META


@app.post("/parse")
async def parse_document(file: UploadFile = File(...)):
    """Step 1: extract raw text from an uploaded PDF or DOCX."""
    content = await file.read()
    text = _extract_text(content, file.filename or "file.txt")
    return {"filename": file.filename, "text": text, "chars": len(text)}


@app.post("/tailor")
async def tailor_cv(
    cv_file:       UploadFile          = File(...),
    tor_text:      str                 = Form(""),
    criteria_json: str                 = Form("[]"),
    project_name:  str                 = Form("Project"),
    bangladesh_project: str            = Form("false"),   # "true" → BD rule active
    sector_groups_json: str            = Form("[]"),      # JSON SectorGroup[]
    deep_analysis: str                 = Form("false"),   # "true" → thorough evidence + inference
):
    """
    Step 2: parse the CV, call Gemini with TOR + matrix context,
    return structured JSON with tailored sections + matrix matches.
    Supports:
      - bangladesh_project=true  → non-Bangladesh experience counts as international
      - sector_groups_json=[{groupLabel, sectors, mode}]  → cumulative/individual sector rules
    """
    # Parse CV — supports PDF/DOCX (incl. tables) and images/scanned CVs via OCR
    cv_bytes = await cv_file.read()
    cv_name = cv_file.filename or "cv.txt"
    if _is_image(cv_name):
        cv_text = await _ocr_image_to_text(cv_bytes, _image_mime(cv_name))
    else:
        cv_text = _extract_text(cv_bytes, cv_name)
    if not cv_text.strip():
        raise HTTPException(400, "Could not extract any text from the CV. If it is a scanned document, upload it as an image (PNG/JPG) so OCR can read it.")

    is_bd_project = bangladesh_project.strip().lower() in ("true", "1", "yes")
    is_deep = deep_analysis.strip().lower() in ("true", "1", "yes")

    try:
        criteria = json.loads(criteria_json)
    except json.JSONDecodeError:
        criteria = []

    try:
        sector_groups = json.loads(sector_groups_json)
    except json.JSONDecodeError:
        sector_groups = []

    # Cache key based on content hash (include rules in hash)
    content_hash = hashlib.sha256(
        (cv_text + tor_text + criteria_json + bangladesh_project + sector_groups_json + deep_analysis).encode()
    ).hexdigest()[:24]
    cache_path = CACHE_DIR / f"{content_hash}.json"
    if cache_path.exists():
        logger.info("Cache hit for %s", content_hash)
        return JSONResponse(json.loads(cache_path.read_text()))

    # Build a prompt if Gemini is available, else mock
    if not GEMINI_API_KEY:
        result = _mock_tailor(cv_text, tor_text, criteria)
    else:
        crit_block = "\n".join(
            f"- {c.get('label', c.get('text', ''))} (max {c.get('maxPoints', c.get('max_score', 2))} pts)"
            for c in criteria[:20]
        )

        # Build special rules block
        special_rules = ""
        if is_bd_project:
            special_rules += (
                "\nBANGLADESH PROJECT RULE: This project runs in Bangladesh. "
                "Any experience the expert gained OUTSIDE Bangladesh counts as international experience "
                "and must be highlighted as such when addressing 'international experience' criteria. "
                "Domestic (Bangladesh) experience is valuable for local context but is NOT international experience.\n"
            )
        if sector_groups:
            rules_lines = []
            for g in sector_groups:
                mode  = g.get("mode", "individual")
                label = g.get("groupLabel", "Sector Group")
                secs  = g.get("sectors", [])
                if mode == "cumulative":
                    rules_lines.append(
                        f"- CUMULATIVE SECTOR GROUP '{label}': sectors [{', '.join(secs)}]. "
                        "Sum the expert's experience years across ALL sectors in this group — a candidate "
                        "with 3 years TVET + 2 years vocational training counts as 5 years total. "
                        "Report the total and which sectors contributed."
                    )
                else:
                    rules_lines.append(
                        f"- INDIVIDUAL SECTOR GROUP '{label}': sectors [{', '.join(secs)}]. "
                        "Score each sector separately. Do not sum across sectors."
                    )
            if rules_lines:
                special_rules += "\nSECTOR EXPERIENCE RULES:\n" + "\n".join(rules_lines) + "\n"

        if is_deep:
            special_rules += (
                "\nDEEP ANALYSIS MODE: Be thorough and exhaustive. For EVERY criterion, actively search the "
                "entire CV for direct AND indirect evidence. If a requirement is not explicitly stated but can "
                "be reasonably inferred from roles, projects, employers, or described activities, credit it and "
                "mark the evidence with '(inferred)'. Categorise the expert's experience by theme and never omit "
                "relevant data — the goal is zero data loss and complete matrix coverage.\n"
            )

        cv_window = 20000 if is_deep else 12000
        prompt = f"""You are a professional CV tailoring expert for development-sector bids (GIZ, EU, UN).
Your task: tailor an expert's CV specifically for the project "{project_name}".
{special_rules}
TERMS OF REFERENCE (excerpt):
\"\"\"
{tor_text[:6000]}
\"\"\"

EVALUATION MATRIX CRITERIA:
{crit_block or "Not provided — infer from the TOR."}

EXPERT CV:
\"\"\"
{cv_text[:cv_window]}
\"\"\"

Instructions:
1. Rewrite the expert's key sections (Professional Summary, Qualifications, Relevant Experience) to directly address the TOR language and evaluation criteria.
2. For each evaluation criterion, cite evidence from the CV as a short phrase (under 25 words). Apply the Bangladesh and sector rules above when present.
3. Compute a tor_match_pct (0-100) based on how well the CV covers the TOR requirements.
4. Return ONLY a valid JSON object. No markdown fences, no prose outside the JSON.

Required JSON schema:
{{
  "expert_name": "string",
  "tor_match_pct": number,
  "sections": [
    {{"section": "string", "original": "string|null", "tailored": "string", "keywords": ["string"]}}
  ],
  "matrix_matches": [
    {{"requirement": "string", "evidence": "string", "score": number, "max_score": number}}
  ],
  "provider": "gemini"
}}"""

        try:
            raw = await _call_gemini(prompt)
            result = _parse_json_loose(raw)
            result.setdefault("provider", "gemini")
        except Exception as e:
            logger.warning("Gemini failed (%s), using mock fallback", e)
            result = _mock_tailor(cv_text, tor_text, criteria)

    cache_path.write_text(json.dumps(result))
    return JSONResponse(result)


@app.post("/generate")
async def generate_docx(request: Request):
    """
    Step 3: receive the tailored JSON + format choice (+ optional person_data),
    generate and stream a .docx or .pdf file.
    person_data is required for template_id == "custom1".
    """
    body = await request.json()
    result      = body.get("result", {})
    template_id = body.get("template_id", "giz")
    person_data = body.get("person_data")   # dict or None
    if template_id not in TEMPLATE_META:
        template_id = "giz"

    expert_name = result.get("expert_name", "Expert").replace(" ", "_")

    if template_id == "latex_modern":
        pdf_bytes = _build_pdf(result, person_data)
        filename = f"{expert_name}_{template_id}_tailored.pdf"
        return StreamingResponse(
            io.BytesIO(pdf_bytes),
            media_type="application/pdf",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'},
        )

    docx_bytes  = _build_docx(result, template_id, person_data)
    filename    = f"{expert_name}_{template_id}_tailored.docx"

    return StreamingResponse(
        io.BytesIO(docx_bytes),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ── Grounded tailor: accepts pre-uploaded Gemini File URIs ────────────────────

class FileRef(BaseModel):
    mime_type: str
    file_uri:  str


class GroundedTailorRequest(BaseModel):
    tor_file:     FileRef
    matrix_file:  FileRef
    cv_file:      FileRef
    project_name: str = "Project"
    template_id:  str = "giz"


@app.post("/grounded-tailor")
async def grounded_tailor(body: GroundedTailorRequest):
    """
    Accept pre-uploaded Gemini File URIs (ToR + Evaluation Matrix + CV).
    This is the NotebookLM-equivalent path: all three documents are grounded
    — the model cannot answer from general knowledge, only the files.
    Returns the same TailorResult JSON structure + a ready-to-download DOCX.
    """
    if not GEMINI_API_KEY:
        raise HTTPException(503, "GEMINI_API_KEY not set — cannot use grounded mode")

    prompt = f"""You are a professional CV tailoring expert for development-sector bids (GIZ, EU, UN).
You have access to THREE attached documents:
  1. The Terms of Reference (ToR)
  2. The Evaluation Matrix (scoring criteria)
  3. The expert's CV

Task: tailor the CV for project "{body.project_name}".
- Rewrite the professional summary to address the TOR language directly
- Identify the strongest evidence for EACH evaluation criterion (cite verbatim, under 25 words)
- Compute tor_match_pct (0–100) based on how well the CV meets the criteria
- ONLY use evidence that genuinely appears in the CV document

Return JSON only (no fences):
{{"expert_name": "string", "tor_match_pct": number,
 "sections": [{{"section": "string", "original": "string|null", "tailored": "string", "keywords": ["string"]}}],
 "matrix_matches": [{{"requirement": "verbatim from TOR", "evidence": "under 25 words from CV", "score": number, "max_score": number}}],
 "provider": "gemini-grounded"}}"""

    file_parts = [
        {"mimeType": body.tor_file.mime_type,    "fileUri": body.tor_file.file_uri},
        {"mimeType": body.matrix_file.mime_type, "fileUri": body.matrix_file.file_uri},
        {"mimeType": body.cv_file.mime_type,     "fileUri": body.cv_file.file_uri},
    ]
    parts = [{"fileData": f} for f in file_parts]
    parts.append({"text": prompt})
    payload = {
        "contents": [{"parts": parts}],
        "generationConfig": {"temperature": 0, "responseMimeType": "application/json"},
    }
    async with httpx.AsyncClient(timeout=180) as client:
        r = await client.post(
            f"https://generativelanguage.googleapis.com/v1beta/models/{GEMINI_MODEL}:generateContent?key={GEMINI_API_KEY}",
            json=payload,
        )
        if not r.is_success:
            raise HTTPException(502, f"Gemini error {r.status_code}: {r.text[:200]}")
        data = r.json()
    raw_text = data["candidates"][0]["content"]["parts"][0]["text"]
    result = _parse_json_loose(raw_text)

    # Cache
    content_hash = hashlib.sha256(
        f"{body.tor_file.file_uri}{body.cv_file.file_uri}".encode()
    ).hexdigest()[:24]
    (CACHE_DIR / f"grounded_{content_hash}.json").write_text(json.dumps(result))

    return JSONResponse(result)


# ── ToR Excerpt extraction ────────────────────────────────────────────────────

@app.post("/extract-tor")
async def extract_tor(file: Optional[UploadFile] = File(None), raw_text: str = Form("")):
    """
    Extract a structured ToR excerpt from EITHER:
      - an uploaded PDF / DOCX / image (scanned docs use Gemini vision OCR), OR
      - pasted raw text (raw_text form field).
    Returns position, summary, required qualifications, key tasks and an excerpt.
    """
    filename = "tor.txt"
    if file is not None and file.filename:
        filename = file.filename
        content = await file.read()
        if _is_image(filename):
            raw_text = await _ocr_image_to_text(content, _image_mime(filename))
        else:
            raw_text = _extract_text(content, filename)
    if not raw_text or not raw_text.strip():
        raise HTTPException(400, "No text found. Upload a readable file/image or paste the ToR text.")

    if not GEMINI_API_KEY:
        # Mock structured excerpt
        result = {
            "position": "Expert",
            "summary": raw_text[:400],
            "required_qualifications": [l.strip() for l in raw_text.split("\n") if len(l.strip()) > 30][:6],
            "key_tasks": [],
            "excerpt_text": raw_text[:4000],
            "provider": "mock",
        }
    else:
        prompt = f"""You are an expert bid manager. From the Terms of Reference (ToR) text below,
extract a clean, structured excerpt that will be used to tailor an expert CV.

TERMS OF REFERENCE:
\"\"\"
{raw_text[:14000]}
\"\"\"

Return ONLY valid JSON (no markdown fences):
{{
  "position": "the expert position / role title",
  "summary": "2-3 sentence summary of the assignment",
  "required_qualifications": ["specific qualification / experience requirement", "..."],
  "key_tasks": ["key task / responsibility", "..."],
  "excerpt_text": "a focused 200-400 word excerpt of the most relevant TOR language for CV tailoring",
  "provider": "gemini"
}}"""
        try:
            raw = await _call_gemini(prompt)
            result = _parse_json_loose(raw)
            result.setdefault("provider", "gemini")
        except Exception as e:
            logger.warning("Gemini ToR extract failed (%s), using mock", e)
            result = {
                "position": "Expert",
                "summary": raw_text[:400],
                "required_qualifications": [l.strip() for l in raw_text.split("\n") if len(l.strip()) > 30][:6],
                "key_tasks": [],
                "excerpt_text": raw_text[:4000],
                "provider": "mock",
            }

    result["raw_text"] = raw_text[:20000]
    result["filename"] = filename
    return JSONResponse(result)


# ── Evaluation Matrix extraction ──────────────────────────────────────────────

@app.post("/extract-matrix")
async def extract_matrix(file: Optional[UploadFile] = File(None), raw_text: str = Form("")):
    """
    Extract role-based evaluation matrices from EITHER a PDF/DOCX/image file OR
    pasted raw text. A single source may define several roles (Team Leader,
    Key Expert 1, ...) — each becomes its own matrix with scored criteria.
    """
    filename = "matrix.txt"
    if file is not None and file.filename:
        filename = file.filename
        content = await file.read()
        if _is_image(filename):
            raw_text = await _ocr_image_to_text(content, _image_mime(filename))
        else:
            raw_text = _extract_text(content, filename)
    if not raw_text or not raw_text.strip():
        raise HTTPException(400, "No text found. Upload a readable file/image or paste the matrix text.")

    if not GEMINI_API_KEY:
        result = {
            "matrices": [
                {
                    "role": "Expert",
                    "criteria": [
                        {"label": l.strip()[:120], "maxPoints": 3}
                        for l in raw_text.split("\n") if len(l.strip()) > 30
                    ][:8],
                }
            ],
            "provider": "mock",
        }
    else:
        prompt = f"""You are an expert bid evaluator. The document below is an evaluation / staffing
matrix that may define scoring criteria for SEVERAL different expert roles.

MATRIX DOCUMENT:
\"\"\"
{raw_text[:16000]}
\"\"\"

Extract ONE evaluation matrix per distinct role. For each role list its scored criteria
with the maximum points available.

Return ONLY valid JSON (no markdown fences):
{{
  "matrices": [
    {{
      "role": "role / position name (e.g. Team Leader)",
      "criteria": [
        {{"label": "the requirement being scored", "maxPoints": number}}
      ]
    }}
  ],
  "provider": "gemini"
}}"""
        try:
            raw = await _call_gemini(prompt)
            result = _parse_json_loose(raw)
            result.setdefault("provider", "gemini")
            if not isinstance(result.get("matrices"), list):
                result = {"matrices": [], "provider": result.get("provider", "gemini")}
        except Exception as e:
            logger.warning("Gemini matrix extract failed (%s), using mock", e)
            result = {
                "matrices": [
                    {
                        "role": "Expert",
                        "criteria": [
                            {"label": l.strip()[:120], "maxPoints": 3}
                            for l in raw_text.split("\n") if len(l.strip()) > 30
                        ][:8],
                    }
                ],
                "provider": "mock",
            }

    result["raw_text"] = raw_text[:20000]
    result["filename"] = filename
    return JSONResponse(result)


# ── CV Creation from Bank (expert data as JSON, no file upload) ───────────────

class ExpertBankInput(BaseModel):
    expert_name: str = ""
    nationality: str = ""
    current_location: str = ""
    level: str = ""
    proposed_position: str = ""
    experience_summary: str = ""
    education_summary: str = ""
    languages_summary: str = ""
    strengths: str = ""
    previous_matrix_matches: list[dict] = []
    tor_text: str = ""
    criteria_json: str = "[]"
    project_name: str = "Project"
    bangladesh_project: bool = False
    sector_groups: list[dict] = []
    deep_analysis: bool = False


@app.post("/tailor-from-json")
async def tailor_from_json(body: ExpertBankInput):
    """CV Creation Wizard: tailor from stored expert bank data (no file upload)."""
    lines = [f"Name: {body.expert_name}"]
    if body.nationality:       lines.append(f"Nationality: {body.nationality}")
    if body.current_location:  lines.append(f"Current Location: {body.current_location}")
    if body.level:             lines.append(f"Level/Role: {body.level}")
    if body.proposed_position: lines.append(f"Proposed Position: {body.proposed_position}")
    if body.education_summary: lines.append(f"\nEDUCATION:\n{body.education_summary}")
    if body.languages_summary: lines.append(f"\nLANGUAGES:\n{body.languages_summary}")
    if body.experience_summary:lines.append(f"\nPROFESSIONAL EXPERIENCE:\n{body.experience_summary}")
    if body.strengths:         lines.append(f"\nKEY STRENGTHS:\n{body.strengths}")
    if body.previous_matrix_matches:
        ev = [f"- {m.get('requirement','')}: {m.get('evidence','')}" for m in body.previous_matrix_matches[:30]]
        lines.append("\nPREVIOUS MATRIX EVIDENCE:\n" + "\n".join(ev))
    cv_text = "\n".join(lines).strip()
    if not cv_text or len(cv_text) < 20:
        raise HTTPException(400, "Expert profile data too sparse. Add more information first.")

    try:
        criteria = json.loads(body.criteria_json)
    except Exception:
        criteria = []

    content_hash = hashlib.sha256(
        (cv_text + body.tor_text + body.criteria_json + body.project_name + str(body.deep_analysis)).encode()
    ).hexdigest()[:24]
    cache_path = CACHE_DIR / f"json_{content_hash}.json"
    if cache_path.exists():
        return JSONResponse(json.loads(cache_path.read_text()))

    if not GEMINI_API_KEY:
        return JSONResponse(_mock_tailor(cv_text, body.tor_text, criteria))

    crit_block = "\n".join(
        f"- {c.get('label', c.get('text', ''))} (max {c.get('maxPoints', c.get('max_score', 2))} pts)"
        for c in criteria[:20]
    )
    special_rules = ""
    if body.bangladesh_project:
        special_rules += "\nBANGLADESH PROJECT RULE: Any experience OUTSIDE Bangladesh counts as international experience.\n"
    for g in body.sector_groups:
        if g.get("mode") == "cumulative":
            special_rules += f"\nCUMULATIVE SECTOR GROUP '{g.get('groupLabel','Group')}': sectors [{', '.join(g.get('sectors',[]))}]. Sum years.\n"
    if body.deep_analysis:
        special_rules += "\nDEEP ANALYSIS: Be exhaustive. Infer competencies from roles/projects (mark '(inferred)'). Zero data loss.\n"

    cv_window = 20000 if body.deep_analysis else 12000
    prompt = f"""You are an expert CV writer for development-sector bids (GIZ, EU, UN).
Task: create a TAILORED CV for expert "{body.expert_name}" proposed as "{body.proposed_position}" for project "{body.project_name}".
{special_rules}
TERMS OF REFERENCE:
\"\"\"
{body.tor_text[:6000]}
\"\"\"
EVALUATION MATRIX:
{crit_block or "Infer from TOR."}
EXPERT PROFILE (use ALL data, lose nothing):
\"\"\"
{cv_text[:cv_window]}
\"\"\"
Write tailored CV sections. Preserve ALL the expert's data — only reframe to match TOR language.
Do NOT invent qualifications. Compute tor_match_pct 0-100.
Return ONLY valid JSON:
{{"expert_name":"string","tor_match_pct":number,"sections":[{{"section":"string","original":"string|null","tailored":"string","keywords":["string"]}}],"matrix_matches":[{{"requirement":"string","evidence":"string","score":number,"max_score":number}}],"provider":"gemini"}}"""

    try:
        raw = await _call_gemini(prompt)
        result = _parse_json_loose(raw)
        result.setdefault("provider", "gemini")
    except Exception as e:
        logger.warning("tailor-from-json Gemini failed (%s), mock fallback", e)
        result = _mock_tailor(cv_text, body.tor_text, criteria)

    cache_path.write_text(json.dumps(result))
    return JSONResponse(result)


@app.post("/compress")
async def compress_docx(file: UploadFile = File(...)):
    """Compress a Word DOCX document by stripping XML revision metadata, optimizing images, and zipping."""
    import zipfile
    import re
    from PIL import Image

    if not file.filename.lower().endswith('.docx'):
        raise HTTPException(400, "Only .docx files are supported for compression.")

    def clean_xml_metadata(xml_data: bytes) -> bytes:
        try:
            text = xml_data.decode("utf-8", errors="ignore")
            # 1. Remove w:rsids block (stores thousands of revision save IDs in settings.xml)
            text = re.sub(r'<w:rsids>.*?</w:rsids>', '', text, flags=re.DOTALL)
            # 2. Strip individual rsid attributes on paragraph/run tags (w:rsidR, w:rsidRPr, etc.)
            text = re.sub(r'\s+w:rsid[A-Za-z0-9]*="[A-Fa-f0-9]+"', '', text)
            # 3. Clean document history, namespaces, and namespaces references
            text = re.sub(r'\s+w:rsidDel="[A-Fa-f0-9]+"', '', text)
            text = re.sub(r'\s+w:rsidP="[A-Fa-f0-9]+"', '', text)
            return text.encode("utf-8")
        except Exception:
            return xml_data

    try:
        in_buf = io.BytesIO(await file.read())
        out_buf = io.BytesIO()

        with zipfile.ZipFile(in_buf, "r") as in_zip:
            with zipfile.ZipFile(out_buf, "w", zipfile.ZIP_DEFLATED, compresslevel=9) as out_zip:
                for item in in_zip.infolist():
                    data = in_zip.read(item.filename)
                    logger.info(f"File inside DOCX: {item.filename} (size: {len(data)} bytes)")
                    
                    # 1. XML Compression: Strip revision logs and RSID attributes
                    if item.filename.lower().endswith(".xml"):
                        data = clean_xml_metadata(data)
                    
                    # 2. Image Compression: Aggressively optimize embedded media
                    elif item.filename.lower().startswith("word/media/"):
                        try:
                            img = Image.open(io.BytesIO(data))
                            
                            # Downscale to 600px width (highly optimized for documents)
                            max_width = 600
                            if img.width > max_width:
                                ratio = max_width / float(img.width)
                                new_height = int(float(img.height) * ratio)
                                img = img.resize((max_width, new_height), Image.Resampling.LANCZOS)

                            img_out = io.BytesIO()
                            
                            # Convert transparency to RGB white background and compress to JPEG format
                            if img.mode in ("RGBA", "LA", "P"):
                                if img.mode == "P":
                                    img = img.convert("RGBA")
                                if "A" in img.getbands():
                                    background = Image.new("RGB", img.size, (255, 255, 255))
                                    background.paste(img, mask=img.split()[-1])
                                    img = background
                                else:
                                    img = img.convert("RGB")
                            elif img.mode != "RGB":
                                img = img.convert("RGB")
                                
                            # Save with quality=45 (extremely high compression with acceptable legibility)
                            img.save(img_out, format="JPEG", quality=45, optimize=True, progressive=True)
                            compressed_data = img_out.getvalue()
                            
                            if len(compressed_data) < len(data):
                                data = compressed_data
                        except Exception as img_err:
                            logger.error(f"Image compress failed for {item.filename}: {img_err}")
                            pass
                    
                    out_zip.writestr(item, data)

        out_buf.seek(0)
        original_size = len(in_buf.getvalue())
        compressed_size = len(out_buf.getvalue())
        ratio = round((1 - compressed_size / original_size) * 100, 1) if original_size else 0
        logger.info(f"SlideSpeak-grade Max-Compress: {original_size} -> {compressed_size} bytes ({ratio}% reduction)")
        
        filename = file.filename
        if not filename.lower().endswith("_compressed.docx"):
            filename = filename.rsplit(".", 1)[0] + "_compressed.docx"

        return StreamingResponse(
            out_buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f'attachment; filename="{filename}"',
                "Access-Control-Expose-Headers": "Content-Disposition"
            }
        )
    except Exception as e:
        logger.error(f"DOCX max compression failed: {e}")
        raise HTTPException(500, f"Failed to compress document: {str(e)}")

